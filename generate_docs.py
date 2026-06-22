from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create Document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ==================== COVER PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('DOKUMENTASI KODE\nE-RAPOR AIRLANGGA')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Sistem Manajemen Pelaporan Akademik')
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run(f'Tanggal: {datetime.now().strftime("%d %B %Y")}').font.size = Pt(12)

doc.add_page_break()

# ==================== DAFTAR ISI ====================
doc.add_heading('DAFTAR ISI', 0)
doc.add_paragraph('1. Pengenalan Sistem', style='List Bullet')
doc.add_paragraph('2. Arsitektur Aplikasi', style='List Bullet')
doc.add_paragraph('3. Data Models', style='List Bullet')
doc.add_paragraph('4. Controllers', style='List Bullet')
doc.add_paragraph('5. Routes & API', style='List Bullet')
doc.add_paragraph('6. User Roles & Permissions', style='List Bullet')

doc.add_page_break()

# ==================== BAB 1: PENGENALAN ====================
doc.add_heading('1. Pengenalan Sistem', level=1)

doc.add_heading('1.1 Gambaran Umum', level=2)
intro_text = """
E-Rapor Airlangga adalah sistem manajemen pelaporan akademik berbasis web yang dibangun dengan Laravel. 
Sistem ini dirancang untuk memudahkan proses pencatatan nilai, penilaian, dan pembuatan rapor siswa. 
Aplikasi ini mendukung multiple roles (Admin, Guru, Orang Tua) dengan fitur-fitur yang disesuaikan 
untuk setiap pengguna.
"""
doc.add_paragraph(intro_text.strip())

doc.add_heading('1.2 Teknologi yang Digunakan', level=2)
tech_list = [
    'Laravel (Framework PHP)',
    'MySQL (Database)',
    'Blade (Template Engine)',
    'Tailwind CSS (Styling)',
    'Vite (Build Tool)',
    'Eloquent ORM (Database Abstraction)'
]
for tech in tech_list:
    doc.add_paragraph(tech, style='List Bullet')

doc.add_heading('1.3 Fitur Utama', level=2)
features = [
    'Manajemen data siswa dan guru',
    'Pencatatan nilai dan penilaian',
    'Pembuatan laporan akademik (rapor)',
    'Tracking kehadiran siswa',
    'Multi-role user management',
    'Dashboard dinamis berdasarkan role',
    'Validasi dan publikasi rapor'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# ==================== BAB 2: ARSITEKTUR ====================
doc.add_heading('2. Arsitektur Aplikasi', level=1)

doc.add_heading('2.1 Struktur Folder', level=2)
structure = """
app/
├── Http/
│   ├── Controllers/     # Business logic untuk request handling
│   ├── Middleware/      # Request validation & authentication
│   └── Requests/        # Form validation rules
├── Models/              # Database models & relationships
├── Mail/                # Email templates & sending logic
└── Providers/           # Service providers for bootstrapping

database/
├── migrations/          # Database schema versioning
└── seeders/             # Initial data seeding

resources/
├── css/                 # Stylesheet files
├── js/                  # JavaScript files
└── views/               # Blade template files

routes/
├── web.php              # Web routes
└── auth.php             # Authentication routes
"""
doc.add_paragraph(structure.strip(), style='List Bullet')

doc.add_heading('2.2 Design Pattern', level=2)
doc.add_paragraph('Model-View-Controller (MVC):', style='List Bullet')
doc.add_paragraph('Separation of concerns antara Models, Controllers, dan Views', style='List Bullet 2')
doc.add_paragraph('Repository Pattern untuk data access abstraction', style='List Bullet')
doc.add_paragraph('Middleware untuk request processing dan authorization', style='List Bullet')

doc.add_page_break()

# ==================== BAB 3: DATA MODELS ====================
doc.add_heading('3. Data Models', level=1)

doc.add_heading('3.1 User Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data pengguna sistem (authentication)', style='List Bullet')
doc.add_paragraph('Tabel Database: users', style='List Bullet')
doc.add_heading('Properties:', level=3)
user_props = ['id', 'name', 'email', 'password', 'role (admin/teacher/parent)', 'email_verified_at', 'created_at', 'updated_at']
for prop in user_props:
    doc.add_paragraph(prop, style='List Bullet')
doc.add_heading('Relations:', level=3)
doc.add_paragraph('hasOne(Teacher) - User sebagai guru', style='List Bullet')
doc.add_paragraph('hasOne(Parents) - User sebagai orang tua', style='List Bullet')
doc.add_heading('Methods Penting:', level=3)
doc.add_paragraph('hasRole($role): Mengecek apakah user memiliki role tertentu', style='List Bullet')

doc.add_heading('3.2 Student Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data siswa', style='List Bullet')
doc.add_paragraph('Tabel Database: students', style='List Bullet')
doc.add_heading('Fillable Fields:', level=3)
student_fields = ['nisn', 'name', 'classroom_id', 'parent_id']
for field in student_fields:
    doc.add_paragraph(field, style='List Bullet')
doc.add_heading('Relations:', level=3)
doc.add_paragraph('belongsTo(Classroom) - Kelas tempat siswa', style='List Bullet')
doc.add_paragraph('belongsTo(Parents) - Orang tua siswa', style='List Bullet')
doc.add_paragraph('hasMany(Grade) - Nilai-nilai siswa', style='List Bullet')
doc.add_paragraph('hasMany(GradeEntry) - Entri penilaian detail', style='List Bullet')
doc.add_paragraph('hasMany(ReportCard) - Rapor siswa', style='List Bullet')
doc.add_paragraph('hasMany(Attendance) - Kehadiran siswa', style='List Bullet')

doc.add_heading('3.3 Teacher Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data guru', style='List Bullet')
doc.add_paragraph('Tabel Database: teachers', style='List Bullet')
doc.add_heading('Fillable Fields:', level=3)
teacher_fields = ['user_id', 'nip', 'gender', 'telp']
for field in teacher_fields:
    doc.add_paragraph(field, style='List Bullet')
doc.add_heading('Relations:', level=3)
doc.add_paragraph('belongsTo(User) - User yang berelasi dengan guru', style='List Bullet')
doc.add_paragraph('hasMany(Grade) - Nilai yang dinilai guru', style='List Bullet')
doc.add_paragraph('hasOne(Classroom) - Kelas yang di-ampu sebagai wali kelas', style='List Bullet')
doc.add_paragraph('hasMany(ClassroomSubjectTeacher) - Pengajaran mata pelajaran', style='List Bullet')

doc.add_heading('3.4 Classroom Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data kelas/rombel', style='List Bullet')
doc.add_paragraph('Tabel Database: classrooms', style='List Bullet')
doc.add_heading('Fillable Fields:', level=3)
classroom_fields = ['name', 'major', 'homeroom_teacher_id']
for field in classroom_fields:
    doc.add_paragraph(field, style='List Bullet')
doc.add_heading('Relations:', level=3)
doc.add_paragraph('belongsTo(Teacher) - Wali kelas', style='List Bullet')
doc.add_paragraph('hasMany(Student) - Daftar siswa di kelas', style='List Bullet')
doc.add_paragraph('hasMany(ClassroomSubjectTeacher) - Mata pelajaran & guru pengampu', style='List Bullet')

doc.add_heading('3.5 Subject Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data mata pelajaran', style='List Bullet')
doc.add_paragraph('Tabel Database: subjects', style='List Bullet')

doc.add_heading('3.6 Grade Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan nilai akhir siswa per mata pelajaran', style='List Bullet')
doc.add_paragraph('Tabel Database: grades', style='List Bullet')
doc.add_heading('Fillable Fields:', level=3)
grade_fields = ['student_id', 'classroom_subject_teacher_id', 'final_score', 'description']
for field in grade_fields:
    doc.add_paragraph(field, style='List Bullet')
doc.add_heading('Casts:', level=3)
doc.add_paragraph('final_score dikonversi menjadi decimal:2', style='List Bullet')

doc.add_heading('3.7 Assessment Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data penilaian/soal', style='List Bullet')
doc.add_paragraph('Tabel Database: assessments', style='List Bullet')
doc.add_heading('Fillable Fields:', level=3)
assessment_fields = ['classroom_subject_teacher_id', 'type', 'name', 'date', 'weight', 'description', 'sequence']
for field in assessment_fields:
    doc.add_paragraph(field, style='List Bullet')
doc.add_heading('Assessment Types:', level=3)
doc.add_paragraph('uh (Ulangan Harian)', style='List Bullet')
doc.add_paragraph('tugas (Tugas)', style='List Bullet')
doc.add_paragraph('pas (Penilaian Akhir Semester)', style='List Bullet')

doc.add_heading('3.8 ReportCard Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data rapor siswa', style='List Bullet')
doc.add_paragraph('Tabel Database: report_cards', style='List Bullet')
doc.add_heading('Fillable Fields:', level=3)
report_fields = ['student_id', 'academic_year_id', 'final_score', 'rank', 'description', 'publish_at', 'is_validated', 'is_submitted', 'status']
for field in report_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('3.9 AcademicYear Model', level=2)
doc.add_paragraph('Keterangan: Model untuk menyimpan data tahun akademik', style='List Bullet')
doc.add_paragraph('Tabel Database: academic_years', style='List Bullet')
doc.add_heading('Fillable Fields:', level=3)
academic_fields = ['year', 'semester', 'is_active']
for field in academic_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('3.10 ClassroomSubjectTeacher Model', level=2)
doc.add_paragraph('Keterangan: Model pivot untuk relasi antara kelas, mata pelajaran, dan guru', style='List Bullet')
doc.add_paragraph('Tabel Database: classroom_subject_teacher', style='List Bullet')

doc.add_page_break()

# ==================== BAB 4: CONTROLLERS ====================
doc.add_heading('4. Controllers', level=1)

doc.add_heading('4.1 AdminDashboardController', level=2)
doc.add_paragraph('Keterangan: Controller untuk dashboard admin', style='List Bullet')
doc.add_heading('Method: index()', level=3)
doc.add_paragraph('Menampilkan statistik: jumlah siswa, guru, orang tua, kelas, dan mata pelajaran', style='List Bullet')
doc.add_paragraph('Mendapatkan tahun akademik yang sedang aktif', style='List Bullet')

doc.add_heading('4.2 StudentController', level=2)
doc.add_paragraph('Keterangan: Controller untuk CRUD data siswa', style='List Bullet')
doc.add_heading('Methods:', level=3)
doc.add_paragraph('index() - Menampilkan daftar siswa dengan pencarian dan paginasi', style='List Bullet')
doc.add_paragraph('create() - Menampilkan form tambah siswa', style='List Bullet')
doc.add_paragraph('store() - Menyimpan data siswa baru (validasi: nisn, name, classroom_id)', style='List Bullet')
doc.add_paragraph('show() - Menampilkan detail siswa', style='List Bullet')
doc.add_paragraph('edit() - Menampilkan form edit siswa', style='List Bullet')
doc.add_paragraph('update() - Mengupdate data siswa', style='List Bullet')

doc.add_heading('4.3 TeacherController', level=2)
doc.add_paragraph('Keterangan: Controller untuk CRUD data guru', style='List Bullet')

doc.add_heading('4.4 ClassroomController', level=2)
doc.add_paragraph('Keterangan: Controller untuk CRUD data kelas', style='List Bullet')

doc.add_heading('4.5 GradesController', level=2)
doc.add_paragraph('Keterangan: Controller untuk manajemen nilai siswa', style='List Bullet')

doc.add_heading('4.6 TeacherDashboardController', level=2)
doc.add_paragraph('Keterangan: Controller untuk dashboard guru', style='List Bullet')

doc.add_heading('4.7 ProfileController', level=2)
doc.add_paragraph('Keterangan: Controller untuk manajemen profil user', style='List Bullet')

doc.add_heading('4.8 SwitchRoleController', level=2)
doc.add_paragraph('Keterangan: Controller untuk switching role pengguna multi-role', style='List Bullet')

doc.add_page_break()

# ==================== BAB 5: ROUTES ====================
doc.add_heading('5. Routes & API', level=1)

doc.add_heading('5.1 Authentication Routes (routes/auth.php)', level=2)
auth_routes = [
    'GET /login - Menampilkan halaman login',
    'POST /login - Memproses login',
    'POST /logout - Memproses logout',
    'GET /forgot-password - Form reset password',
    'POST /forgot-password - Kirim link reset',
    'GET /reset-password/{token} - Form reset password',
    'POST /reset-password - Update password baru',
    'GET /verify-email - Halaman verifikasi email',
    'GET /verify-email/{id}/{hash} - Verifikasi email',
    'POST /activate - Aktivasi akun baru'
]
for route in auth_routes:
    doc.add_paragraph(route, style='List Bullet')

doc.add_heading('5.2 Public Routes (routes/web.php)', level=2)
public_routes = [
    'GET / - Halaman welcome',
    'GET /dashboard - Redirect ke dashboard sesuai role',
    'GET /activate/{token} - Form aktivasi akun'
]
for route in public_routes:
    doc.add_paragraph(route, style='List Bullet')

doc.add_heading('5.3 Admin Routes', level=2)
admin_routes = [
    'GET /admin/dashboard - Dashboard admin',
    'GET /admin/students - Daftar siswa',
    'POST /admin/students - Tambah siswa',
    'GET /admin/students/{id}/edit - Edit siswa',
    'PUT /admin/students/{id} - Update siswa',
    'GET /admin/teachers - Daftar guru',
    'GET /admin/classrooms - Daftar kelas',
    'GET /admin/subjects - Daftar mata pelajaran',
    'GET /admin/academic-years - Daftar tahun akademik'
]
for route in admin_routes:
    doc.add_paragraph(route, style='List Bullet')

doc.add_heading('5.4 Teacher Routes', level=2)
teacher_routes = [
    'GET /teacher/dashboard - Dashboard guru',
    'GET /teacher/kelas-saya - Kelas yang diampu',
    'GET /teacher/grades - Manajemen nilai',
    'GET /teacher/attendance - Kehadiran siswa'
]
for route in teacher_routes:
    doc.add_paragraph(route, style='List Bullet')

doc.add_heading('5.5 Parent Routes', level=2)
parent_routes = [
    'GET /parent/dashboard - Dashboard orang tua',
    'GET /parent/children - Daftar anak',
    'GET /parent/report-cards - Lihat rapor anak'
]
for route in parent_routes:
    doc.add_paragraph(route, style='List Bullet')

doc.add_page_break()

# ==================== BAB 6: USER ROLES ====================
doc.add_heading('6. User Roles & Permissions', level=1)

doc.add_heading('6.1 Role: Admin', level=2)
doc.add_paragraph('Deskripsi: Administrator sistem dengan akses penuh', style='List Bullet')
doc.add_heading('Permissions:', level=3)
admin_perms = [
    'Manajemen data siswa (CRUD)',
    'Manajemen data guru (CRUD)',
    'Manajemen data kelas (CRUD)',
    'Manajemen mata pelajaran',
    'Manajemen tahun akademik',
    'Validasi rapor siswa',
    'Lihat dashboard admin dengan statistik'
]
for perm in admin_perms:
    doc.add_paragraph(perm, style='List Bullet')

doc.add_heading('6.2 Role: Teacher (Guru)', level=2)
doc.add_paragraph('Deskripsi: Guru yang mengampu kelas dan mata pelajaran', style='List Bullet')
doc.add_heading('Permissions:', level=3)
teacher_perms = [
    'Lihat daftar siswa di kelas',
    'Input nilai penilaian siswa',
    'Lihat dashboard guru',
    'Lihat kelas yang diampu',
    'Kelola penilaian (assessment)',
    'Input kehadiran siswa',
    'Lihat rapor siswa'
]
for perm in teacher_perms:
    doc.add_paragraph(perm, style='List Bullet')

doc.add_heading('6.3 Role: Parent (Orang Tua)', level=2)
doc.add_paragraph('Deskripsi: Orang tua/wali siswa untuk monitoring akademik', style='List Bullet')
doc.add_heading('Permissions:', level=3)
parent_perms = [
    'Lihat data diri anaknya',
    'Lihat nilai dan rapor anak',
    'Lihat kehadiran anak',
    'Lihat perkembangan akademik anak'
]
for perm in parent_perms:
    doc.add_paragraph(perm, style='List Bullet')

doc.add_page_break()

# ==================== FOOT NOTE ====================
doc.add_heading('Catatan Tambahan', level=1)
doc.add_paragraph(
    'Dokumentasi ini mencakup overview dari sistem E-Rapor Airlangga. '
    'Untuk informasi lebih detail, silakan merujuk ke kode source di repository. '
    'Sistem ini dirancang untuk kemudahan maintenance dan scalability dengan '
    'mengikuti Laravel best practices dan SOLID principles.'
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('End of Documentation')
footer_run.font.italic = True
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Save Document
output_path = 'DOKUMENTASI_E-RAPOR_AIRLANGGA.docx'
doc.save(output_path)
print(f'✓ Dokumentasi berhasil dibuat: {output_path}')
